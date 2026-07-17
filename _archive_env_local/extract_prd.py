from docx import Document

doc = Document('C:/Users/fll/Downloads/Demo PRD-审小助 (1).docx')
with open('prd_text.txt', 'w', encoding='utf-8') as f:
    f.write('=== PRD 内容提取 ===\n')
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ''
        f.write(f'[{style}] {text}\n')

    for ti, table in enumerate(doc.tables):
        f.write(f'\n=== Table {ti} ===\n')
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            f.write(' | '.join(cells) + '\n')
print('已写入 prd_text.txt')
